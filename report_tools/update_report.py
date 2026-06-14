import docx
from pathlib import Path

ROOT_DIR = Path(__file__).parent.parent

def replace_paragraph_text(p, old_text, new_text):
    if old_text not in p.text:
        return False
    replaced = False
    for run in p.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)
            replaced = True
    if not replaced:
        p.text = p.text.replace(old_text, new_text)
    return True

def update_report():
    doc_path = ROOT_DIR / "RAPOR.docx"
    doc = docx.Document(doc_path)
    
    # ---------------------------------------------------------------------------
    # 1. Update Tables
    # ---------------------------------------------------------------------------
    
    # Table 1: Test Results
    doc.tables[1].rows[1].cells[3].text = "Achieved 73.0% on German and 73.8% on HELOC."
    doc.tables[1].rows[3].cells[3].text = "100% validity achieved for all tested cases."
    doc.tables[1].rows[5].cells[3].text = "Differences computed successfully (0.105 for gender)."
    
    # Table 2: Performance Evaluation
    doc.tables[2].rows[1].cells[2].text = "F1: 0.806 (German) / 0.722 (HELOC), ROC-AUC: 0.766 (German) / 0.809 (HELOC)."
    doc.tables[2].rows[3].cells[2].text = "SHAP: 1.85 (German) / 1.23 (HELOC), LIME: 2.83 (German) / 1.31 (HELOC), IG: 3.34 (German) / 2.76 (HELOC)."
    
    # Table 3: Dataset Results
    # German Credit
    doc.tables[3].rows[1].cells[1].text = "0.730"
    doc.tables[3].rows[1].cells[2].text = "0.812"
    doc.tables[3].rows[1].cells[3].text = "0.800"
    doc.tables[3].rows[1].cells[4].text = "0.806"
    doc.tables[3].rows[1].cells[5].text = "0.766"
    doc.tables[3].rows[1].cells[6].text = "0.390"
    # HELOC
    doc.tables[3].rows[2].cells[1].text = "0.738"
    doc.tables[3].rows[2].cells[2].text = "0.734"
    doc.tables[3].rows[2].cells[3].text = "0.711"
    doc.tables[3].rows[2].cells[4].text = "0.722"
    doc.tables[3].rows[2].cells[5].text = "0.809"
    doc.tables[3].rows[2].cells[6].text = "0.500"
    
    # Table 4: Recourse Results
    # German Credit project_search
    doc.tables[4].rows[1].cells[2].text = "1.000"
    doc.tables[4].rows[1].cells[3].text = "12"
    doc.tables[4].rows[1].cells[4].text = "3.250"
    doc.tables[4].rows[1].cells[5].text = "3.000"
    doc.tables[4].rows[1].cells[6].text = "0.204"
    doc.tables[4].rows[1].cells[7].text = "0.594"
    # HELOC project_search
    doc.tables[4].rows[3].cells[2].text = "1.000"
    doc.tables[4].rows[3].cells[3].text = "12"
    doc.tables[4].rows[3].cells[4].text = "4.667"
    doc.tables[4].rows[3].cells[5].text = "5.000"
    doc.tables[4].rows[3].cells[6].text = "0.150"
    doc.tables[4].rows[3].cells[7].text = "0.650"
    
    # Table 5: Details OpenXAI Methods
    # Update Column 1 "occlusion" to "ig" for German Credit
    doc.tables[5].rows[1].cells[1].text = "ig"
    # German Credit ig
    doc.tables[5].rows[1].cells[2].text = "0.325"
    doc.tables[5].rows[1].cells[3].text = "0.097"
    doc.tables[5].rows[1].cells[4].text = "3.34"
    doc.tables[5].rows[1].cells[5].text = "0.661"
    doc.tables[5].rows[1].cells[6].text = "0.914"
    doc.tables[5].rows[1].cells[7].text = "0.081"
    
    # German Credit shap
    doc.tables[5].rows[2].cells[2].text = "0.448"
    doc.tables[5].rows[2].cells[3].text = "0.242"
    doc.tables[5].rows[2].cells[4].text = "1.85"
    doc.tables[5].rows[2].cells[5].text = "1.179"
    doc.tables[5].rows[2].cells[6].text = "1.075"
    doc.tables[5].rows[2].cells[7].text = "0.074"
    
    # German Credit lime
    doc.tables[5].rows[3].cells[2].text = "0.314"
    doc.tables[5].rows[3].cells[3].text = "0.111"
    doc.tables[5].rows[3].cells[4].text = "2.83"
    doc.tables[5].rows[3].cells[5].text = "2.871"
    doc.tables[5].rows[3].cells[6].text = "3.138"
    doc.tables[5].rows[3].cells[7].text = "0.066"
    
    # Update Column 1 "occlusion" to "ig" for HELOC
    doc.tables[5].rows[4].cells[1].text = "ig"
    # HELOC ig
    doc.tables[5].rows[4].cells[2].text = "0.467"
    doc.tables[5].rows[4].cells[3].text = "0.169"
    doc.tables[5].rows[4].cells[4].text = "2.76"
    doc.tables[5].rows[4].cells[5].text = "1.416"
    doc.tables[5].rows[4].cells[6].text = "0.870"
    doc.tables[5].rows[4].cells[7].text = "-"
    
    # HELOC shap
    doc.tables[5].rows[5].cells[2].text = "0.526"
    doc.tables[5].rows[5].cells[3].text = "0.426"
    doc.tables[5].rows[5].cells[4].text = "1.23"
    doc.tables[5].rows[5].cells[5].text = "1.361"
    doc.tables[5].rows[5].cells[6].text = "0.623"
    doc.tables[5].rows[5].cells[7].text = "-"
    
    # HELOC lime
    doc.tables[5].rows[6].cells[2].text = "0.403"
    doc.tables[5].rows[6].cells[3].text = "0.306"
    doc.tables[5].rows[6].cells[4].text = "1.31"
    doc.tables[5].rows[6].cells[5].text = "2.197"
    doc.tables[5].rows[6].cells[6].text = "1.660"
    doc.tables[5].rows[6].cells[7].text = "-"

    # ---------------------------------------------------------------------------
    # 2. Update Paragraphs
    # ---------------------------------------------------------------------------
    
    p220_text = "\tThe predictive module is implemented in PyTorch as a Multi-Layer Perceptron (MLP) within the CreditDNN class. The network architecture consists of linear layers, Rectified Linear Unit (ReLU) activations, and dropout layers configured to mitigate overfitting. For both the German Credit and FICO HELOC datasets, the hidden layers are configured as [64, 32] hidden structures. Mathematically, the forward propagation is defined as a sequence of linear layers, batch normalization, ReLU activation, and dropout, followed by a final linear layer outputting raw logits. The models are trained using Binary Cross-Entropy Loss with the AdamW optimizer (learning rate of 0.0005 and batch size of 64 for German Credit, and learning rate of 0.0007 and batch size of 128 for HELOC). To optimize classification decisions, we select the decision threshold on the validation split by maximizing the F1-score, yielding a threshold of 0.390 for German Credit and 0.500 for HELOC. If the Sigmoid output probability is equal to or greater than the threshold, the application is rejected."
    p221_text = "The explainability module wraps three local attribution explainers to interpret the frozen PyTorch model. The SHAP wrapper instantiates a shap.KernelExplainer using a background training set of 50 samples to compute local Shapley feature attributions. The LIME wrapper utilizes the lime.lime_tabular.LimeTabularExplainer to evaluate raw feature values. The Integrated Gradients (IG) explainer is implemented as a path-based method that computes the gradients of the model's output with respect to the input along a linear path from a baseline (training set mean) to the input instance, integrated over 50 steps: IG_i(x) = (x_i - x_baseline_i) * mean(grad_i(x_baseline + alpha * (x - x_baseline))). OpenXAI evaluations are performed on the test set, computing Prediction Gap on Important features (PGI) and Relative Input Stability (RIS) by introducing random Gaussian noise to the scaled inputs and recording attribution shifts."
    p282_text = "\tA key finding in our quantitative explainer evaluation is that the Relative Input Stability (RIS) and Relative Representation Stability (RRS) metrics are reported on a natural-log stability scale, where values closer to 0 indicate high stability (with a ratio closer to 1.0). Mathematically, these metrics are formulated using a natural logarithm over the maximum ratio of the Euclidean norm of the explanation attribution delta to the norm of the perturbation delta. In our results, Integrated Gradients (IG) demonstrates the highest stability for German Credit, achieving an RIS of 0.661 and an RRS of 0.099. This represents Lipschitz-like smoothness and local explanation robustness, indicating that when small, random noise is introduced to an applicant's profile, the IG explainer does not exhibit volatile, unpredictable shifts in feature importance rankings. Conversely, LIME exhibits higher sensitivity to input noise, achieving an RIS of 2.871 on German Credit."
    p283_text = "To determine which explanation method is superior, we must evaluate them across different dimensions, as no single explainer dominates on all metrics. Under the faithfulness dimension, measured by the Prediction Gap on Important features (PGI), SHAP (KernelExplainer) is superior, achieving the highest scores of 0.448 on German Credit and 0.526 on HELOC. Conversely, under the feature separation dimension, Integrated Gradients (IG) attribution is superior, achieving low PGU scores (0.097 on German Credit and 0.169 on HELOC) and the highest PGI/PGU ratios. Under the stability dimension, IG is superior on German Credit (RIS of 0.661), whereas SHAP is superior on HELOC (RIS of 1.361), representing robust local approximations under input perturbations."

    for p in doc.paragraphs:
        if p.text.strip().startswith("The predictive module is implemented in PyTorch as a Multi-Layer Perceptron (MLP) within the CreditDNN class."):
            p.text = p220_text
        elif p.text.strip().startswith("The explainability module wraps three local attribution explainers to interpret the frozen PyTorch model."):
            p.text = p221_text
        elif p.text.strip().startswith("A critical finding in our quantitative explainer evaluation is the presence of negative values"):
            p.text = p282_text
        elif p.text.strip().startswith("To determine which explanation method is superior, we must evaluate them across different dimensions"):
            p.text = p283_text

    replacements = [
        ("SHAP, LIME, and Occlusion", "SHAP, LIME, and Integrated Gradients (IG)"),
        ("SHAP, LIME, Occlusion", "SHAP, LIME, Integrated Gradients (IG)"),
        ("SHAP, LIME, and Occlusion attribution", "SHAP, LIME, and Integrated Gradients (IG) attribution"),
        ("SHAP, LIME, and Occlusion explainers", "SHAP, LIME, and Integrated Gradients (IG) explainers"),
        ("SHAP, LIME and Occlusion", "SHAP, LIME and Integrated Gradients (IG)"),
        ("explainers (SHAP, LIME, and Occlusion) using OpenXAI metrics", "explainers (SHAP, LIME, and Integrated Gradients (IG)) using OpenXAI metrics"),
        ("Occlusion attribution (frequently referred to as feature ablation or local sensitivity analysis)", 
         "Integrated Gradients (IG) attribution (a path-based method that computes gradients along a path from a baseline to the input)"),
        ("SHAP, LIME, and Occlusion, highlighting the trade-offs", "SHAP, LIME, and Integrated Gradients (IG), highlighting the trade-offs"),
        ("SHAP, LIME, and Occlusion wrappers", "SHAP, LIME, and Integrated Gradients (IG) wrappers"),
        ("custom occlusion attribution loops", "integrated gradients path integration loops"),
        ("averaging a 0.80 ROC-AUC across the datasets.", "averaging a 0.79 ROC-AUC across the datasets."),
        ("yielding an ROC-AUC of 0.840 for the Statlog German Credit dataset and 0.808 for the FICO HELOC dataset. The high recall values (0.689 for German Credit and 0.800 for HELOC) confirm",
         "yielding an ROC-AUC of 0.766 for the Statlog German Credit dataset and 0.809 for the FICO HELOC dataset. The recall values (0.800 for German Credit and 0.711 for HELOC) confirm"),
        ("averaging only 1.55 changes for German Credit", "averaging only 3.25 changes for German Credit"),
        ("LIME is the recommended choice. Using LIME prevents erratic, jumping visualizations in the UI under minor slider adjustments, providing a smooth user experience. Finally, Occlusion attribution should be used as a fast, computationally inexpensive diagnostic baseline during initial model feature selection to quickly isolate and discard irrelevant variables.",
         "LIME is the recommended choice. Using LIME prevents erratic, jumping visualizations in the UI under minor slider adjustments, providing a smooth user experience. Finally, Integrated Gradients (IG) attribution should be used as a path-based baseline during initial model feature selection to quickly isolate and discard irrelevant variables."),
        ("For the gender attribute, the model exhibited a demographic parity difference of 0.667 and an equalized odds difference of 1.0. These high disparity values stem from deep-seated gender imbalances in historical credit registries, where female applicants were severely underrepresented and faced higher default labels. For the age attribute, the model demonstrated a demographic parity difference of 0.084 and an equalized odds difference of 0.073.",
         "For the gender attribute, the model exhibited a demographic parity difference of 0.105 and an equalized odds difference of 0.105. These disparity values stem from gender imbalances in historical credit registries, where female applicants were underrepresented. For the age attribute, the model demonstrated a demographic parity difference of 0.124 and an equalized odds difference of 0.164."),
        ("Occlusion offers the cleanest separation of important features.", "Integrated Gradients (IG) offers the cleanest separation of important features."),
        ("The UCI Statlog German Credit dataset, consisting of 1,000 instances, was divided into a training subset of 699 samples (489 good and 210 bad risk), a validation subset of 150 samples (105 good and 45 bad risk), and a held-out test subset of 151 samples (106 good and 45 bad risk). The FICO HELOC dataset, comprising 9,871 instances, was split into a training subset of 6,909 samples (3,314 good and 3,595 bad risk), a validation subset of 1,480 samples (710 good and 770 bad risk), and a test subset of 1,482 samples (711 good and 771 bad risk).",
         "The UCI Statlog German Credit dataset, consisting of 1,000 instances, was divided into a training subset of 640 samples, a validation subset of 160 samples, and a held-out test subset of 200 samples. The FICO HELOC dataset, comprising 9,872 cleaned instances, was split into a training subset of 6,319 samples, a validation subset of 1,579 samples, and a test subset of 1,974 samples."),
        ("optimal configuration selected consists of a feed-forward network with two hidden layers containing 64 and 32 neurons respectively, trained with a batch size of 64, a learning rate of 0.001, L2 weight decay of 0.0001, a dropout rate of 0.25, and early stopping patience of 30 epochs, reaching its best validation loss at epoch 11.",
         "optimal configuration selected consists of a feed-forward network with two hidden layers containing 64 and 32 neurons respectively, trained with a batch size of 64, a learning rate of 0.0005, L2 weight decay of 0.0001, a dropout rate of 0.30, and early stopping patience of 20 epochs, reaching its best validation loss at epoch 11."),
        ("For the FICO HELOC dataset, the optimal network consists of two hidden layers containing 128 and 64 neurons, trained with a batch size of 64, a learning rate of 0.0005, L2 weight decay of 0.0001, a dropout rate of 0.20, and early stopping patience of 30 epochs, reaching its best validation loss at epoch 9.",
         "For the FICO HELOC dataset, the optimal network consists of two hidden layers containing 64 and 32 neurons, trained with a batch size of 128, a learning rate of 0.0007, L2 weight decay of 0.0001, a dropout rate of 0.25, and early stopping patience of 12 epochs."),
        ("For the German Credit model, the optimal threshold was selected at 0.595, yielding a validation accuracy of 75.3% and a validation F1 score of 62.6%, which translated to a test accuracy of 80.1%, test precision of 66.0%, test recall of 68.9%, test F1 score of 67.4%, and a ROC-AUC of 84.0%. The confusion matrix on the held-out test set resulted in 90 true negatives, 16 false positives, 14 false negatives, and 31 true positives. For the HELOC model, the optimal threshold was calibrated at 0.415, yielding a validation accuracy of 75.2% and a validation F1 score of 77.8%, which resulted in a test accuracy of 73.2%, test precision of 71.7%, test recall of 80.0%, test F1 score of 75.7%, and a ROC-AUC of 80.8%. The test confusion matrix on the HELOC test set resulted in 468 true negatives, 243 false positives, 154 false negatives, and 617 true positives.",
         "For the German Credit model, the optimal threshold was selected at 0.390, yielding a test accuracy of 73.0%, test precision of 81.2%, test recall of 80.0%, test F1 score of 80.6%, and a ROC-AUC of 76.6%. The confusion matrix on the held-out test set resulted in 34 true negatives, 26 false positives, 28 false negatives, and 112 true positives. For the HELOC model, the optimal threshold was calibrated at 0.500, yielding a test accuracy of 73.8%, test precision of 73.4%, test recall of 71.1%, test F1 score of 72.2%, and a ROC-AUC of 80.9%. The test confusion matrix on the HELOC test set resulted in 784 true negatives, 244 false positives, 274 false negatives, and 673 true positives."),
    ]
    
    for p in doc.paragraphs:
        for old_text, new_text in replacements:
            replace_paragraph_text(p, old_text, new_text)
            
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for old_text, new_text in replacements:
                        replace_paragraph_text(p, old_text, new_text)
                        
    doc.save(doc_path)
    print("Report updated successfully!")

if __name__ == "__main__":
    update_report()
