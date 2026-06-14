import docx
from pathlib import Path
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

ROOT_DIR = Path(__file__).parent.parent

def update_report():
    doc_path = ROOT_DIR / "RAPOR.docx"
    doc = docx.Document(doc_path)
    
    # ---------------------------------------------------------------------------
    # 1. Update Section 3.2 System Design/Architecture paragraph
    # ---------------------------------------------------------------------------
    target_p197_start = "The final system architecture is structured as a five-stage sequential pipeline"
    p197_found = False
    for p in doc.paragraphs:
        if p.text.strip().startswith(target_p197_start):
            old_phrase = "using our custom categorical-aware search and the official dice-ml library"
            new_phrase = "using a custom PyTorch gradient-descent DiCE optimization framework (optimizing for model target, proximity, diversity, and categorical constraints) and the official dice-ml library"
            if old_phrase in p.text:
                p.text = p.text.replace(old_phrase, new_phrase)
                p197_found = True
                print("Updated Section 3.2 architecture paragraph.")
                break
    if not p197_found:
        print("Warning: Could not find or update Section 3.2 architecture paragraph.")
    
    # ---------------------------------------------------------------------------
    # 2. Update Section 4.2 Detailed Implementation recourse paragraph
    # ---------------------------------------------------------------------------
    target_p222_start = "The counterfactual recourse module implements a custom categorical-aware heuristic search"
    p222_found = False
    new_p222_text = (
        "\tThe counterfactual recourse module implements a custom gradient-descent-based "
        "Diverse Counterfactual Explanations (DiCE) optimization framework in PyTorch. Unlike "
        "traditional heuristic searches, our custom DiCE implementation defines recourse generation "
        "as a multi-objective optimization problem. Specifically, for a rejected applicant x, we "
        "optimize a pool of candidate counterfactuals C = {c^(1), c^(2), ..., c^(k)} by minimizing "
        "a loss function defined as:"
    )
    new_p222_formula = (
        "L_total = w_pred * L_BCE(f(c), y*) + w_prox * (||c - x||_1 / d) - w_div * Diversity(C) + w_cat * L_cat(c)"
    )
    new_p222_explanation = (
        "where w_pred, w_prox, w_div, and w_cat are weights for the prediction target, proximity (L1 norm), "
        "pairwise diversity, and categorical consistency respectively. The prediction loss L_BCE ensures "
        "that the counterfactual candidates flip the model output to the desired approved class y* (1 = Good). "
        "The proximity loss minimizes the changes relative to the original profile x, scaled by feature dimension d. "
        "The diversity loss promotes different recourse paths by maximizing pairwise L1 distances among candidates. "
        "The categorical regularization loss L_cat enforces that one-hot encoded dummy variables are binary and sum to 1. "
        "To ensure feasibility, we apply a projection step at each training iteration, clamping continuous features "
        "within their empirical bounds and resetting immutable features (e.g., age, gender) to their original values."
    )
    
    # Let's find the paragraph and replace it, and insert formula and explanation paragraphs after it
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith(target_p222_start):
            # Update current paragraph
            p.text = new_p222_text
            
            # Insert formula paragraph
            f_el = OxmlElement('w:p')
            p._element.addnext(f_el)
            f_p = Paragraph(f_el, p._parent)
            f_p.text = new_p222_formula
            f_p.style = doc.styles['Normal']
            
            # Insert explanation paragraph
            exp_el = OxmlElement('w:p')
            f_p._element.addnext(exp_el)
            exp_p = Paragraph(exp_el, f_p._parent)
            exp_p.text = new_p222_explanation
            exp_p.style = doc.styles['Normal']
            
            p222_found = True
            print("Updated Section 4.2 recourse description and inserted formula.")
            break
            
    if not p222_found:
        print("Warning: Could not find or update Section 4.2 recourse paragraph.")

    # ---------------------------------------------------------------------------
    # 3. Update Appendix A description and code block
    # ---------------------------------------------------------------------------
    desc_start = "The second snippet showcases the candidate generation heuristic"
    new_desc_text = (
        "The second snippet showcases the custom gradient descent DiCE optimization loop implemented in our "
        "recourse module in PyTorch, including the pairwise diversity and categorical regularization loss functions."
    )
    desc_found = False
    for p in doc.paragraphs:
        if p.text.strip().startswith(desc_start):
            p.text = new_desc_text
            desc_found = True
            print("Updated Appendix A code description.")
            break
    if not desc_found:
        print("Warning: Could not find or update Appendix A description.")

    # Now let's replace the code paragraphs from `def make_candidate_pool(` up to `return candidates`
    start_idx = -1
    end_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if "def make_candidate_pool(" in p.text:
            start_idx = i
        if start_idx != -1 and "return candidates" in p.text:
            end_idx = i
            break
            
    if start_idx != -1 and end_idx != -1:
        print(f"Found old make_candidate_pool code block from paragraph {start_idx} to {end_idx}.")
        
        new_code_lines = [
            "def pairwise_diversity(candidates: torch.Tensor) -> torch.Tensor:",
            "    if candidates.shape[0] <= 1:",
            "        return torch.tensor(0.0)",
            "",
            "    distances = torch.cdist(candidates, candidates, p=1) / candidates.shape[1]",
            "    mask = torch.triu(torch.ones_like(distances), diagonal=1).bool()",
            "    return distances[mask].mean()",
            "",
            "def categorical_regularization(",
            "    candidates: torch.Tensor,",
            "    categorical_groups: list[list[int]],",
            ") -> torch.Tensor:",
            "    if not categorical_groups:",
            "        return torch.tensor(0.0)",
            "",
            "    losses = []",
            "    for group in categorical_groups:",
            "        values = candidates[:, group]",
            "        sum_loss = (values.sum(dim=1) - 1.0).pow(2).mean()",
            "        binary_loss = (values * (1.0 - values)).abs().mean()",
            "        losses.append(sum_loss + binary_loss)",
            "    return torch.stack(losses).mean()",
            "",
            "def project_candidates(",
            "    candidates: torch.Tensor,",
            "    original: torch.Tensor,",
            "    dataset: LoadedDataset,",
            ") -> None:",
            "    candidates.clamp_(dataset.lower_bounds, dataset.upper_bounds)",
            "    if dataset.immutable_indices:",
            "        candidates[:, dataset.immutable_indices] = original[:, dataset.immutable_indices]",
            "",
            "    decreasing_features = (",
            "        HELOC_DECREASING_FEATURES",
            "        if dataset.name == 'heloc'",
            "        else GERMAN_DECREASING_NUMERIC_FEATURES",
            "    )",
            "    for feature in decreasing_features:",
            "        if feature in dataset.feature_names:",
            "            index = dataset.feature_names.index(feature)",
            "            candidates[:, index] = torch.minimum(candidates[:, index], original[:, index])",
            "",
            "    for group in dataset.categorical_groups:",
            "        candidates[:, group].clamp_(0.0, 1.0)",
            "        group_sum = candidates[:, group].sum(dim=1, keepdim=True).clamp_min(1e-6)",
            "        candidates[:, group] = candidates[:, group] / group_sum",
            "",
            "    for group in dataset.immutable_groups:",
            "        candidates[:, group] = original[:, group]",
            "",
            "def optimize_counterfactuals(",
            "    dataset: LoadedDataset,",
            "    query_encoded: np.ndarray,",
            "    desired_class: int,",
            "    config: CounterfactualConfig,",
            ") -> np.ndarray:",
            "    original = torch.from_numpy(query_encoded.astype(np.float32)).unsqueeze(0)",
            "    initial = original.repeat(config.total_cfs, 1)",
            "    noise = torch.randn_like(initial) * 0.08",
            "    candidates = nn.Parameter(initial + noise)",
            "",
            "    target = torch.full((config.total_cfs,), float(desired_class), dtype=torch.float32)",
            "    optimizer = torch.optim.Adam([candidates], lr=config.learning_rate)",
            "    loss_fn = nn.BCEWithLogitsLoss()",
            "",
            "    with torch.no_grad():",
            "        project_candidates(candidates, original, dataset)",
            "",
            "    for _ in range(config.steps):",
            "        optimizer.zero_grad()",
            "        logits = dataset.model(candidates)",
            "        prediction_loss = loss_fn(logits, target)",
            "        proximity_loss = torch.abs(candidates - original).mean()",
            "        diversity_loss = pairwise_diversity(candidates)",
            "        category_loss = categorical_regularization(candidates, dataset.categorical_groups)",
            "        loss = (",
            "            config.prediction_weight * prediction_loss",
            "            + config.proximity_weight * proximity_loss",
            "            - config.diversity_weight * diversity_loss",
            "            + config.categorical_weight * category_loss",
            "        )",
            "        loss.backward()",
            "        optimizer.step()",
            "        with torch.no_grad():",
            "            project_candidates(candidates, original, dataset)",
            "",
            "    return candidates.detach().cpu().numpy()"
        ]
        
        num_old_paras = end_idx - start_idx + 1
        num_new_lines = len(new_code_lines)
        
        # Overwrite existing paragraphs as much as possible
        for offset in range(min(num_old_paras, num_new_lines)):
            p_idx = start_idx + offset
            doc.paragraphs[p_idx].text = new_code_lines[offset]
            
        if num_new_lines > num_old_paras:
            # Insert paragraphs after the last updated paragraph
            current_p = doc.paragraphs[start_idx + num_old_paras - 1]
            for offset in range(num_old_paras, num_new_lines):
                new_p_el = OxmlElement('w:p')
                current_p._element.addnext(new_p_el)
                new_p = Paragraph(new_p_el, current_p._parent)
                new_p.text = new_code_lines[offset]
                new_p.style = doc.styles['Normal']
                current_p = new_p
        elif num_new_lines < num_old_paras:
            # Delete remaining paragraphs
            paras_to_delete = []
            for offset in range(num_new_lines, num_old_paras):
                paras_to_delete.append(doc.paragraphs[start_idx + offset])
            for p in paras_to_delete:
                p._element.getparent().remove(p._element)
                
        print(f"Successfully replaced old code block with {num_new_lines} new lines.")
    else:
        print("Warning: Could not find make_candidate_pool code block in report.")

    # Save document
    doc.save(doc_path)
    print("Report saved successfully!")

if __name__ == "__main__":
    update_report()
