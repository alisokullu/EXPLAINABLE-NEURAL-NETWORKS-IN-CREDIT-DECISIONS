import os
import docx

def main():
    doc_path = r'c:/Users/aliso/Downloads/xai-credit-decision-main/4992FinalReport1011165.docx'
    doc = docx.Document(doc_path)
    
    idx_a = None
    for idx, p in enumerate(doc.paragraphs):
        if p.style.name == 'Heading 1' and p.text.strip().startswith('APPENDIX A:'):
            idx_a = idx
            print(f"Found APPENDIX A at idx {idx}")
            break
            
    if idx_a is None:
        print("Error: APPENDIX A heading not found.")
        return
        
    # Paragraph at idx_a + 1 is the GitHub repository intro.
    # We delete starting from paragraph at idx_a + 2 until we hit APPENDIX B.
    delete_idx = idx_a + 2
    deleted_count = 0
    
    while delete_idx < len(doc.paragraphs):
        p = doc.paragraphs[delete_idx]
        if p.style.name == 'Heading 1' and p.text.strip().startswith('APPENDIX B:'):
            print(f"Reached APPENDIX B heading at current idx {delete_idx}. Stopping deletion.")
            break
            
        p_element = p._p
        parent = p_element.getparent()
        parent.remove(p_element)
        deleted_count += 1
        
    print(f"Deleted {deleted_count} paragraphs from Appendix A.")
    doc.save(doc_path)
    print("Document saved successfully!")

if __name__ == '__main__':
    main()
