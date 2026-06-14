import docx
from docx.shared import Pt

def format_normal(p):
    p.style = 'Normal'
    for r in p.runs:
        if not r.font.name:
            r.font.name = 'Times New Roman'
        if not r.font.size:
            r.font.size = Pt(12)

def replace_in_runs(p, old, new):
    """Replace text across runs carefully."""
    full_text = p.text
    if old not in full_text:
        return False
    # Simple approach: if single run contains it
    for r in p.runs:
        if old in r.text:
            r.text = r.text.replace(old, new)
            return True
    # Cross-run approach: rebuild paragraph text
    new_full = full_text.replace(old, new)
    # Clear all runs and set first run text
    if p.runs:
        p.runs[0].text = new_full
        for r in p.runs[1:]:
            r.text = ''
    return True

def main():
    doc_path = r'c:/Users/aliso/Downloads/xai-credit-decision-main/4992FinalReport1011165.docx'
    doc = docx.Document(doc_path)

    fixes_applied = []

    for idx, p in enumerate(doc.paragraphs):
        text = p.text
        if not text.strip():
            continue

        changed = False

        # ---------------------------------------------------------------
        # FIX 1: "Dashboard Fariness Page" -> "Dashboard Fairness Page"
        # (paragraphs and table-of-figures entries)
        # ---------------------------------------------------------------
        if 'Fariness' in text:
            for r in p.runs:
                if 'Fariness' in r.text:
                    r.text = r.text.replace('Fariness', 'Fairness')
                    changed = True
            if changed:
                fixes_applied.append(f"P{idx}: Fixed 'Fariness' -> 'Fairness'")

        text = p.text  # refresh

        # ---------------------------------------------------------------
        # FIX 2: "adillik karar destek paneli" -> "Fairness & Decision Support Panel"
        # ---------------------------------------------------------------
        if 'adillik karar destek paneli' in text.lower():
            old = 'adillik karar destek paneli (Fairness & Decision Support Panel)'
            new = 'Fairness & Decision Support Panel'
            if old in text:
                if replace_in_runs(p, old, new):
                    fixes_applied.append(f"P{idx}: Fixed 'adillik karar destek paneli'")
                    changed = True
            else:
                # try without parenthetical
                old2 = 'adillik karar destek paneli'
                if replace_in_runs(p, old2, 'Fairness & Decision Support Panel'):
                    fixes_applied.append(f"P{idx}: Fixed 'adillik karar destek paneli' (variant)")
                    changed = True

        text = p.text  # refresh

        # ---------------------------------------------------------------
        # FIX 3: "adillik calibration tools" -> "fairness calibration tools"
        # ---------------------------------------------------------------
        if 'adillik calibration' in text.lower():
            for r in p.runs:
                if 'adillik calibration' in r.text.lower():
                    r.text = r.text.replace('adillik calibration tools', 'fairness calibration tools')
                    r.text = r.text.replace('Adillik calibration tools', 'fairness calibration tools')
                    changed = True
            if changed:
                fixes_applied.append(f"P{idx}: Fixed 'adillik calibration tools'")

        text = p.text  # refresh

        # ---------------------------------------------------------------
        # FIX 4: Citation swap [4]<->[5] in Literature Review paragraphs
        # P204: Agarwal et al. [5] should be [4]
        # P205: Mothilal et al. [4] should be [5]
        # Also fix any other occurrences in body
        # ---------------------------------------------------------------
        # P204: OpenXAI - Agarwal cited as [5], should be [4]
        if idx == 204 and 'Agarwal et al. [5]' in text:
            if replace_in_runs(p, 'Agarwal et al. [5]', 'Agarwal et al. [4]'):
                fixes_applied.append(f"P{idx}: Fixed Agarwal [5] -> [4]")
                changed = True

        # P205: DiCE/Mothilal cited as [4], should be [5]
        if idx == 205 and 'Mothilal et al. [4]' in text:
            if replace_in_runs(p, 'Mothilal et al. [4]', 'Mothilal et al. [5]'):
                fixes_applied.append(f"P{idx}: Fixed Mothilal [4] -> [5]")
                changed = True

        # P219 tech stack: DiCE [5] is correct (Mothilal/DiCE), SHAP [3], LIME [2] are correct
        # OpenXAI is [4] - check if any OpenXAI [5] in tech stack
        if 'OpenXAI' in text and '[5]' in text and idx != 364:
            if replace_in_runs(p, 'OpenXAI [5]', 'OpenXAI [4]'):
                fixes_applied.append(f"P{idx}: Fixed OpenXAI [5] -> [4]")
                changed = True

        # ---------------------------------------------------------------
        # FIX 5: LIME stability contradiction in P290
        # Original: "Using LIME prevents erratic, jumping visualizations in the UI 
        #  under minor slider adjustments, providing a smooth user experience."
        # Problem: Table 5 shows LIME RIS=2.871 (worst stability). 
        # Fix: Clarify that LIME is recommended for UI not because it's more stable,
        # but because it's faster/more responsive and provides human-readable 
        # surrogate explanations. Remove the "prevents erratic" false claim.
        # ---------------------------------------------------------------
        if idx == 290:
            old_lime = ("Using LIME prevents erratic, jumping visualizations in the UI under minor slider adjustments, "
                        "providing a smooth user experience.")
            new_lime = ("Although LIME exhibits higher input sensitivity (RIS of 2.871 on German Credit) compared to IG and SHAP, "
                        "its surrogate model architecture produces human-readable, locally-faithful linear approximations that are "
                        "more interpretable for non-technical risk analysts. Furthermore, LIME's computationally lightweight "
                        "surrogate fitting enables faster real-time recalculations when analysts adjust applicant profiles in "
                        "the dashboard, making it the preferred choice for interactive decision-support interfaces despite its "
                        "lower formal stability score.")
            if old_lime in text:
                if replace_in_runs(p, old_lime, new_lime):
                    fixes_applied.append(f"P{idx}: Fixed LIME stability contradiction in recommendation paragraph")
                    changed = True
            else:
                # Try partial match
                old_partial = "Using LIME prevents erratic, jumping visualizations"
                if old_partial in text:
                    # Find and replace broader section
                    new_full_text = text.replace(
                        "However, for live user-facing dashboard interfaces and what-if simulation tools where risk analysts dynamically modify applicant profiles, LIME is the recommended choice. Using LIME prevents erratic, jumping visualizations in the UI under minor slider adjustments, providing a smooth user experience.",
                        "However, for live user-facing dashboard interfaces and what-if simulation tools where risk analysts dynamically modify applicant profiles, LIME remains a practical choice. Although LIME exhibits higher input sensitivity (RIS of 2.871 on German Credit) compared to IG and SHAP, its surrogate model architecture produces human-readable, locally-faithful linear approximations that are more interpretable for non-technical risk analysts. Furthermore, LIME's computationally lightweight surrogate fitting enables faster real-time recalculations when analysts adjust applicant profiles in the dashboard, making it the preferred choice for interactive decision-support interfaces despite its lower formal stability score."
                    )
                    if new_full_text != text and p.runs:
                        p.runs[0].text = new_full_text
                        for r in p.runs[1:]:
                            r.text = ''
                        fixes_applied.append(f"P{idx}: Fixed LIME stability contradiction (partial match)")
                        changed = True

    # ---------------------------------------------------------------
    # Also fix in tables (for adillik/fariness)
    # ---------------------------------------------------------------
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    cell_text = para.text
                    if 'Fariness' in cell_text:
                        for run in para.runs:
                            if 'Fariness' in run.text:
                                run.text = run.text.replace('Fariness', 'Fairness')
                                fixes_applied.append(f"Table{t_idx} R{r_idx} C{c_idx}: Fixed 'Fariness'")
                    if 'adillik' in cell_text.lower():
                        for run in para.runs:
                            if 'adillik' in run.text.lower():
                                run.text = (run.text
                                    .replace('adillik karar destek paneli (Fairness & Decision Support Panel)', 'Fairness & Decision Support Panel')
                                    .replace('adillik karar destek paneli', 'Fairness & Decision Support Panel')
                                    .replace('adillik calibration tools', 'fairness calibration tools')
                                    .replace('Adillik calibration tools', 'fairness calibration tools'))
                                fixes_applied.append(f"Table{t_idx} R{r_idx} C{c_idx}: Fixed 'adillik'")

    doc.save(doc_path)

    print("=== ALL FIXES APPLIED ===")
    for fix in fixes_applied:
        print(f"  OK: {fix}")
    print(f"\nTotal: {len(fixes_applied)} fix(es) applied. Document saved.")

if __name__ == '__main__':
    main()
