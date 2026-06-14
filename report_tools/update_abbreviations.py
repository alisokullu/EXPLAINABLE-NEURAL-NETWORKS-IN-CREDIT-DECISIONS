import docx
from pathlib import Path
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

ROOT_DIR = Path(__file__).parent.parent
DOC_PATH = ROOT_DIR.parent / "4992FinalReport1011165.docx"

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._edge = paragraph._parent = None

def update_abbreviations():
    if not DOC_PATH.exists():
        print(f"Error: {DOC_PATH} does not exist!")
        return

    doc = docx.Document(DOC_PATH)
    
    # 1. Locate LIST OF ABBREVIATIONS and 1. Introduction
    start_idx = -1
    end_idx = -1
    
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "LIST OF ABBREVIATIONS":
            start_idx = i
        elif start_idx != -1 and p.text.strip() == "1. Introduction":
            end_idx = i
            break
            
    if start_idx == -1 or end_idx == -1:
        print(f"Error: Could not locate abbreviation boundaries (start: {start_idx}, end: {end_idx})")
        return
        
    print(f"Found abbreviation list between paragraph {start_idx} and {end_idx}.")
    
    # 2. Collect paragraphs to delete
    to_delete = doc.paragraphs[start_idx + 1 : end_idx]
    print(f"Deleting {len(to_delete)} paragraphs between boundaries...")
    for p in to_delete:
        delete_paragraph(p)
        
    # Re-load the document structure to get updated paragraph references
    doc.save(DOC_PATH)
    doc = docx.Document(DOC_PATH)
    
    # Find start paragraph again in refreshed document
    start_para = None
    for p in doc.paragraphs:
        if p.text.strip() == "LIST OF ABBREVIATIONS":
            start_para = p
            break
            
    if not start_para:
        print("Error: Refreshed document missing 'LIST OF ABBREVIATIONS'!")
        return

    # List of new abbreviations in alphabetical order
    abbreviations = [
        ("BCE", "Binary Cross-Entropy"),
        ("DICE", "Diverse Counterfactual Explanations"),
        ("DNN", "Deep Neural Network"),
        ("ECOA", "Equal Credit Opportunity Act"),
        ("GDPR", "General Data Protection Regulation"),
        ("IG", "Integrated Gradients"),
        ("LIME", "Local Interpretable Model-agnostic Explanations"),
        ("MLP", "Multi-Layer Perceptron"),
        ("PGI", "Prediction Gap on Important features"),
        ("PGU", "Prediction Gap on Unimportant features"),
        ("RIS", "Relative Input Stability"),
        ("ROC-AUC", "Receiver Operating Characteristic - Area Under the Curve"),
        ("ROS", "Relative Output Stability"),
        ("RRS", "Relative Representation Stability"),
        ("SHAP", "Shapley Additive Explanations"),
        ("UCI", "University of California, Irvine"),
        ("XAI", "Explainable Artificial Intelligence"),
    ]
    
    # Insert new abbreviation paragraphs
    current_p = start_para
    for abbrev, desc in abbreviations:
        new_p_el = OxmlElement('w:p')
        current_p._element.addnext(new_p_el)
        new_p = Paragraph(new_p_el, current_p._parent)
        new_p.style = doc.styles['Normal (Web)']
        
        # Abbreviation in Bold
        r_abbrev = new_p.add_run(abbrev)
        r_abbrev.bold = True
        r_abbrev.font.name = "Times New Roman"
        r_abbrev.font.size = Pt(12)
        
        # Tab Separator and Description in Regular
        # For longer abbreviations, we use one tab, for shorter ones we use two to align nicely
        sep = "\t\t" if len(abbrev) < 6 else "\t"
        r_desc = new_p.add_run(sep + desc)
        r_desc.bold = False
        r_desc.font.name = "Times New Roman"
        r_desc.font.size = Pt(12)
        
        current_p = new_p
        
    # Save modified document
    doc.save(DOC_PATH)
    print("Abbreviations list updated successfully!")

if __name__ == "__main__":
    update_abbreviations()
