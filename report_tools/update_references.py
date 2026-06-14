import os
import docx
from docx.shared import Pt

def insert_paragraph_after_paragraph(doc, p, text, style='Normal'):
    new_p = doc.add_paragraph()
    new_p.style = style
    p_element = p._p
    parent = p_element.getparent()
    parent.insert(parent.index(p_element) + 1, new_p._p)
    return new_p

def main():
    doc_path = r'c:/Users/aliso/Downloads/xai-credit-decision-main/4992FinalReport1011165.docx'
    doc = docx.Document(doc_path)
    
    idx_207 = None
    idx_210 = None
    idx_ref_7 = None
    
    # Locate target paragraphs
    for idx, p in enumerate(doc.paragraphs):
        text_strip = p.text.strip()
        if text_strip.startswith("The technical stack is built entirely on the Python ecosystem"):
            idx_207 = idx
            print(f"Found Tech Stack paragraph at idx {idx}")
        elif text_strip.startswith("The empirical validation of our framework is conducted on two distinct public credit datasets"):
            idx_210 = idx
            print(f"Found Datasets paragraph at idx {idx}")
        elif text_strip.startswith("[7]") and "FICO" in text_strip:
            idx_ref_7 = idx
            print(f"Found Reference [7] at idx {idx}")
            
    if idx_207 is None or idx_210 is None or idx_ref_7 is None:
        print("Error: Could not locate all target paragraphs.")
        return
        
    p207 = doc.paragraphs[idx_207]
    p210 = doc.paragraphs[idx_210]
    p_ref_7 = doc.paragraphs[idx_ref_7]
        
    # 1. Update Section 3.3 (Tech Stack) with citation [8]
    # We replace within the first run (Run 0)
    orig_text_207 = p207.runs[0].text
    if "Microsoft's Fairlearn library [8]" not in orig_text_207:
        p207.runs[0].text = orig_text_207.replace("Microsoft's Fairlearn library", "Microsoft's Fairlearn library [8]")
        print("Updated Tech Stack with citation [8].")
    else:
        print("Tech Stack already has citation [8].")
        
    # 2. Update Section 3.4 (Datasets) with citations [6] and [7]
    # We replace within Run 2
    orig_text_210_r2 = p210.runs[2].text
    updated_210 = False
    
    if "sourced from the UCI Machine Learning Repository [6]" not in orig_text_210_r2:
        orig_text_210_r2 = orig_text_210_r2.replace(
            "sourced from the UCI Machine Learning Repository", 
            "sourced from the UCI Machine Learning Repository [6]"
        )
        updated_210 = True
        
    if "the FICO HELOC dataset [7], containing" not in orig_text_210_r2:
        orig_text_210_r2 = orig_text_210_r2.replace(
            "the FICO HELOC dataset, containing", 
            "the FICO HELOC dataset [7], containing"
        )
        updated_210 = True
        
    if updated_210:
        p210.runs[2].text = orig_text_210_r2
        print("Updated Datasets paragraph with citations [6] and [7].")
    else:
        print("Datasets paragraph already has citations [6] and [7].")
        
    # 3. Add Reference [8] after Reference [7]
    next_p_idx = idx_ref_7 + 1
    next_p = doc.paragraphs[next_p_idx]
    
    if next_p.text.strip().startswith("[8]"):
        print("Reference [8] already exists.")
    else:
        ref_p = insert_paragraph_after_paragraph(doc, p_ref_7, "", style='Normal')
        ref_p.paragraph_format.space_before = Pt(0)
        ref_p.paragraph_format.space_after = Pt(6)
        ref_p.paragraph_format.line_spacing = 1.15
        
        r1 = ref_p.add_run("[8]")
        r1.font.name = 'Calibri'
        
        r2 = ref_p.add_run('\xa0S. Bird et al., "Fairlearn: A toolkit for assessing and improving fairness in AI," Microsoft Technical Report MSR-TR-2020-32, 2020.')
        r2.font.name = 'Calibri'
        print("Reference [8] inserted successfully.")
        
    doc.save(doc_path)
    print("Document saved successfully!")

if __name__ == '__main__':
    main()
