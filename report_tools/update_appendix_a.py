import os
import docx
from docx.shared import Pt

FILES_AND_DIRS = [
    ("dashboard.py", "Streamlit-based interactive web dashboard. It hosts the graphical interface for model performance overview, global explanations (SHAP/LIME), case-level inspections (feature attributions, Integrated Gradients), and interactive counterfactual recourse simulation."),
    ("train_dnn.py", "Handles training, validation, early stopping, and hyperparameter tuning of the PyTorch Deep Neural Network (DNN) model on the FICO HELOC dataset."),
    ("train_german_dnn.py", "Performs the equivalent training, threshold tuning, and model evaluation pipelines for the German Credit dataset, handling categorical/numerical column transformations."),
    ("clean_datasets.py", "Command-line script that runs the data cleaning and preprocessing pipelines on raw data sources and saves the processed datasets."),
    ("evaluate_xai.py", "Quantitative XAI evaluation script. It benchmarks SHAP and LIME local explanations against OpenXAI faithfulness (PGI/PGU) and stability (RIS/ROS/RRS) metrics."),
    ("generate_shap.py", "Computes and serializes precalculated global SHAP explanations for both HELOC and German Credit models to enable fast loading in the dashboard."),
    ("generate_lime.py", "Computes and serializes precalculated global LIME local explanations for both models."),
    ("generate_ig.py", "Computes global and local Integrated Gradients explanations for the models."),
    ("generate_counterfactuals.py", "Implements the custom actionability-aware counterfactual recourse generation framework (using DiCE) for rejected loan applicants."),
    ("src/", "Contains core project modules, including src/data_cleaning.py, which implements dataset cleaning, missing value imputation, and class-stratified splits."),
    ("dataSets/", "Stores the original raw datasets, including the FICO HELOC dataset and the UCI German Credit raw database."),
    ("cleanedDataSets/", "Contains cleaned datasets ready for training, alongside the data cleaning report log in JSON format."),
    ("artifacts/", "Holds all precomputed models and explanation results:"),
    ("  └─ dnn_german/ & dnn_heloc/", "Trained weights (model.pt), validation metrics, and preprocessor pipeline states."),
    ("  └─ shap/ & lime/ & ig/", "Serialized global and local explanation attributions."),
    ("  └─ dice/", "DiCE counterfactual configurations, recourse recommendations, and full recourse tables."),
    ("  └─ xai_metrics/", "Faithfulness and stability evaluation summary logs in JSON format."),
    ("report_tools/", "Contains auxiliary Python scripts used for document updates, references, abbreviations, and table formatting."),
    ("test_*.py", "Automated unit tests validating model integrity, explanation shape contracts, and dashboard response consistency."),
    ("requirements.txt", "Defines the required library dependencies (PyTorch, Streamlit, SHAP, LIME, DiCE, Fairlearn, etc.) to set up and run the codebase."),
    ("pytest.ini", "Configuration settings for running tests via the pytest framework.")
]

def update_docx(file_path):
    if not os.path.exists(file_path):
        print(f"Skipping: {file_path} (File not found)")
        return
        
    print(f"Updating: {file_path}")
    try:
        doc = docx.Document(file_path)
    except Exception as e:
        print(f"Error opening {file_path}: {e}")
        return
        
    idx_a = None
    for idx, p in enumerate(doc.paragraphs):
        if p.style.name == 'Heading 1' and p.text.strip().startswith('APPENDIX A:'):
            idx_a = idx
            break
            
    if idx_a is None:
        print("Error: APPENDIX A heading not found.")
        return
        
    # Paragraph after heading is the intro text. We update the URL in it.
    intro_p = doc.paragraphs[idx_a + 1]
    old_intro_text = intro_p.text
    new_url = "https://github.com/alisokullu/EXPLAINABLE-NEURAL-NETWORKS-IN-CREDIT-DECISIONS"
    
    # Simple search and replace for the github URL
    if "https://github.com/canberkylcn/xai-credit-decision" in old_intro_text:
        intro_p.text = old_intro_text.replace("https://github.com/canberkylcn/xai-credit-decision", new_url)
    elif "github.com" in old_intro_text:
        # Fallback replacement if the URL was slightly different
        import re
        intro_p.text = re.sub(r'https://github\.com/\S+', new_url, old_intro_text)
    else:
        intro_p.text = f"The complete software repository for the explainable neural network project is hosted publicly on GitHub at: {new_url}. The repository is organized into distinct directories containing the raw data split logs, trained PyTorch model checkpoints, evaluation results, and modular Python scripts. The core architecture is designed to be fully reproducible, implementing leakage-safe data splits, automated hyperparameter tuning, local explanation calculations using SHAP and LIME, and actionable counterfactual recourse generation."
        
    print("Updated GitHub repository URL in intro paragraph.")
    
    # We clear out any empty paragraphs after the intro paragraph up until APPENDIX B.
    delete_idx = idx_a + 2
    to_delete = []
    
    while delete_idx < len(doc.paragraphs):
        p = doc.paragraphs[delete_idx]
        if p.style.name == 'Heading 1' and p.text.strip().startswith('APPENDIX B:'):
            break
        to_delete.append(p)
        delete_idx += 1
        
    # Delete them
    for p in to_delete:
        p_element = p._p
        parent = p_element.getparent()
        parent.remove(p_element)
        
    print(f"Cleaned {len(to_delete)} leftover paragraphs.")
    
    # Now we insert our new file descriptions directly below the intro paragraph.
    # The parent index where we insert is parent.index(intro_p._p)
    parent = intro_p._p.getparent()
    current_insert_idx = parent.index(intro_p._p) + 1
    
    # Add a transition paragraph
    transition_p = doc.add_paragraph()
    transition_p.style = 'Normal'
    transition_p.paragraph_format.space_before = Pt(6)
    transition_p.paragraph_format.space_after = Pt(12)
    transition_p.paragraph_format.line_spacing = 1.15
    run = transition_p.add_run("The detailed structure and descriptions of the files and directories in the repository are provided below:")
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    
    parent.insert(current_insert_idx, transition_p._p)
    current_insert_idx += 1
    
    # Add files/dirs
    for item_name, description in FILES_AND_DIRS:
        p = doc.add_paragraph()
        p.style = 'Normal'
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        
        # In docx, a bullet list can be bulleted by setting paragraph style to 'List Bullet' or adding a bullet symbol
        # Let's add a bullet character manually so it matches standard spacing and does not depend on custom template list styles.
        bullet_run = p.add_run("•  ")
        bullet_run.font.name = 'Calibri'
        bullet_run.font.size = Pt(11)
        
        bold_run = p.add_run(item_name)
        bold_run.bold = True
        bold_run.font.name = 'Calibri'
        bold_run.font.size = Pt(11)
        
        colon_run = p.add_run(": ")
        colon_run.font.name = 'Calibri'
        colon_run.font.size = Pt(11)
        
        desc_run = p.add_run(description)
        desc_run.font.name = 'Calibri'
        desc_run.font.size = Pt(11)
        
        parent.insert(current_insert_idx, p._p)
        current_insert_idx += 1
        
    # Save the document
    try:
        doc.save(file_path)
        print("Successfully updated and saved the document!")
    except PermissionError:
        print(f"Error: {file_path} is currently locked or open. Please close Microsoft Word and try again.")
    except Exception as e:
        print(f"Error saving {file_path}: {e}")

def main():
    paths = [
        r'c:/Users/aliso/Downloads/xai-credit-decision-main/4992FinalReport1011165.docx',
        r'c:/Users/aliso/Downloads/xai-credit-decision-main/4992FinalReport1011165_updated.docx'
    ]
    for path in paths:
        update_docx(path)

if __name__ == '__main__':
    main()
