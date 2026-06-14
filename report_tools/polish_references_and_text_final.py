import os
import docx
from docx.shared import Pt

def add_reference_entry(doc, idx_app_a_p, number, authors, title, journal_or_conf, rest):
    p = idx_app_a_p.insert_paragraph_before()
    p.style = 'Normal'
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    
    r_num = p.add_run(f"[{number}]\xa0")
    r_num.font.name = 'Calibri'
    r_num.font.size = Pt(11)
    
    r_auth = p.add_run(authors)
    r_auth.font.name = 'Calibri'
    r_auth.font.size = Pt(11)
    
    if title:
        r_title = p.add_run(title)
        r_title.font.name = 'Calibri'
        r_title.font.size = Pt(11)
        
    if journal_or_conf:
        r_jc = p.add_run(journal_or_conf)
        r_jc.font.name = 'Calibri'
        r_jc.font.size = Pt(11)
        r_jc.italic = True
        
    r_rest = p.add_run(rest)
    r_rest.font.name = 'Calibri'
    r_rest.font.size = Pt(11)
    return p

def format_normal_paragraph(p):
    p.style = 'Normal'
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    for r in p.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)

def main():
    doc_path = r'c:/Users/aliso/Downloads/xai-credit-decision-main/4992FinalReport1011165.docx'
    doc = docx.Document(doc_path)
    
    # We will locate target paragraphs dynamically by prefix/contains matching.
    p_50 = None
    p_162 = None
    p_206 = None
    p_209 = None
    p_232 = None
    p_267 = None
    p_286 = None
    p_321 = None
    p_335 = None
    p_ref_head = None
    p_ref_1 = None
    p_app_a = None
    
    for idx, p in enumerate(doc.paragraphs):
        text_strip = p.text.strip()
        
        # Paragraph 50 (Executive Summary / Intro)
        if text_strip.startswith("The developed end-to-end local pipeline was rigorously validated"):
            p_50 = p
            print(f"Found P50 at index {idx}")
            
        # Paragraph 162 (Introduction / Proposed Method)
        elif text_strip.startswith("The primary objective of this project is to develop, evaluate, and implement"):
            p_162 = p
            print(f"Found P162 at index {idx}")
            
        # Paragraph 206 (Tech Stack / Section 3.3)
        elif text_strip.startswith("The technical stack is built entirely on the Python ecosystem"):
            p_206 = p
            print(f"Found P206 at index {idx}")
            
        # Paragraph 209 (Datasets / Section 3.4)
        elif text_strip.startswith("The empirical validation of our framework is conducted on two distinct public credit datasets"):
            p_209 = p
            print(f"Found P209 at index {idx}")
            
        # Paragraph 232 (System Assembly / Section 3.6)
        elif text_strip.startswith("System assembly is achieved by integrating the modular Python classes"):
            p_232 = p
            print(f"Found P232 at index {idx}")
            
        # Paragraph 267 (Section 6.1 Results)
        elif text_strip.startswith("Fourth, regarding algorithmic fairness and mitigation on the German Credit dataset"):
            p_267 = p
            print(f"Found P267 at index {idx}")
            
        # Paragraph 286 (Section 7 Conclusion / Summary)
        elif text_strip.startswith("All original and modified objectives established during the project proposal stage"):
            p_286 = p
            print(f"Found P286 at index {idx}")
            
        # Paragraph 321 (Section 8 Ethical Considerations)
        elif text_strip.startswith("In credit scoring, data-driven systems run the risk of reinforcing historical systemic inequalities"):
            p_321 = p
            print(f"Found P321 at index {idx}")
            
        # Paragraph 335 (Section 9.3 Future Work)
        elif text_strip.startswith("There are several promising avenues for future research to build upon the established framework"):
            p_335 = p
            print(f"Found P335 at index {idx}")
            
        # References heading
        elif p.style.name.startswith("Heading") and text_strip == "References":
            p_ref_head = p
            print(f"Found References Heading at index {idx}")
            
        # First reference entry
        elif text_strip.startswith("[1]") and p_ref_1 is None:
            p_ref_1 = p
            print(f"Found First Reference [1] at index {idx}")
            
        # Appendix A heading
        elif p.style.name == 'Heading 1' and text_strip.startswith("APPENDIX A:"):
            p_app_a = p
            print(f"Found APPENDIX A Heading at index {idx}")

    # Check that we found everything we need
    missing = []
    if p_50 is None: missing.append("p_50")
    if p_162 is None: missing.append("p_162")
    if p_206 is None: missing.append("p_206")
    if p_209 is None: missing.append("p_209")
    if p_232 is None: missing.append("p_232")
    if p_267 is None: missing.append("p_267")
    if p_286 is None: missing.append("p_286")
    if p_321 is None: missing.append("p_321")
    if p_335 is None: missing.append("p_335")
    if p_ref_head is None: missing.append("p_ref_head")
    if p_ref_1 is None: missing.append("p_ref_1")
    if p_app_a is None: missing.append("p_app_a")
    
    if missing:
        print(f"Error: Could not locate all target paragraphs. Missing: {missing}")
        return

    # --- 1. Modify Paragraph 50 (Executive Summary / Intro) ---
    txt = p_50.text
    txt = txt.replace("SHAP, LIME, and Integrated Gradients (IG)", "SHAP [3], LIME [2], and Integrated Gradients (IG)")
    txt = txt.replace("DiCE for counterfactual recourse generation,", "DiCE [5] for counterfactual recourse generation,")
    txt = txt.replace("Fairlearn for demographic bias auditing.", "Fairlearn [6] for demographic bias auditing.")
    p_50.text = txt
    format_normal_paragraph(p_50)
    print("Updated P50 with in-text citations.")

    # --- 2. Modify Paragraph 162 (Introduction / Proposed Method) ---
    txt = p_162.text
    txt = txt.replace("LIME and SHAP", "LIME [2] and SHAP [3]")
    txt = txt.replace("DiCE for counterfactual recourse recommendations,", "DiCE [5] for counterfactual recourse recommendations,")
    txt = txt.replace("Fairlearn for demographic bias auditing.", "Fairlearn [6] for demographic bias auditing.")
    p_162.text = txt
    format_normal_paragraph(p_162)
    print("Updated P162 with in-text citations.")

    # --- 3. Modify Paragraph 206 (Tech Stack / Section 3.3) ---
    txt = p_206.text
    # Add citations for shap and lime libraries and dice-ml, and change Fairlearn citation
    txt = txt.replace("shap and lime libraries", "SHAP [3] and LIME [2] libraries")
    txt = txt.replace("dice-ml library", "DiCE [5] library")
    txt = txt.replace("Microsoft's Fairlearn library [8]", "Microsoft's Fairlearn library [6]")
    p_206.text = txt
    format_normal_paragraph(p_206)
    print("Updated P206 with in-text citations.")

    # --- 4. Modify Paragraph 209 (Datasets / Section 3.4) ---
    txt = p_209.text
    txt = txt.replace("UCI Machine Learning Repository [6]", "UCI Machine Learning Repository [7]")
    txt = txt.replace("FICO HELOC dataset [7]", "FICO HELOC dataset [8]")
    p_209.text = txt
    format_normal_paragraph(p_209)
    print("Updated P209 with in-text citations.")

    # --- 5. Modify Paragraph 232 (System Assembly / Section 3.6) ---
    txt = p_232.text
    txt = txt.replace("SHAP/LIME", "SHAP [3]/LIME [2]")
    txt = txt.replace("Fairlearn dynamic threshold", "Fairlearn [6] dynamic threshold")
    p_232.text = txt
    format_normal_paragraph(p_232)
    print("Updated P232 with in-text citations.")

    # --- 6. Modify Paragraph 267 (Section 6.1 Results) ---
    p_267.text = ("Fourth, regarding algorithmic fairness and mitigation on the German Credit dataset, "
                  "the baseline model exhibited a demographic parity difference of 10.48% (with a baseline approval rate "
                  "of 61.67% for female applicants and 72.14% for male applicants) and an equalized odds difference of 10.48% "
                  "for the gender attribute, and a demographic parity difference of 12.42% for the age attribute (binarized at 30 years). "
                  "By applying Fairlearn's [6] ThresholdOptimizer (post-processing match), we successfully mitigated these disparities, "
                  "reducing the gender demographic parity difference from 10.48% to 1.67% (with optimized approval rates "
                  "of 83.33% for female applicants and 85.00% for male applicants). From a Management Engineering perspective, "
                  "this fairness calibration is evaluated against model performance; in our test splits, the group-specific "
                  "threshold adjustment resulted in a model accuracy of 75.50% compared to the baseline 73.00%, demonstrating "
                  "that threshold tuning can align fairness criteria without deteriorating overall credit classification validity. "
                  "These results have been integrated into our Fairness & Decision Support panel, showing that group-specific "
                  "threshold adjustments can enforce equity and regulatory compliance in Automated Decision-Making (ADM) systems.")
    format_normal_paragraph(p_267)
    print("Updated P267 (Section 6.1 Results) with exact numerical metrics and Fairlearn [6] citation.")

    # --- 7. Modify Paragraph 286 (Section 7 Conclusion / Summary) ---
    txt = p_286.text
    txt = txt.replace("conducted using Fairlearn to audit", "conducted using Fairlearn [6] to audit")
    p_286.text = txt
    format_normal_paragraph(p_286)
    print("Updated P286 with in-text citations.")

    # --- 8. Modify Paragraph 321 (Section 8 Ethical Considerations) ---
    txt = p_321.text
    txt = txt.replace("historical systemic inequalities. Our project", "historical systemic inequalities [9]. Our project")
    txt = txt.replace("using Microsoft's Fairlearn library", "using Microsoft's Fairlearn library [6]")
    txt = txt.replace("demographic parity difference of 0.105 and an equalized odds difference of 0.105", 
                      "demographic parity difference of 10.48% (0.1048) and an equalized odds difference of 10.48% (0.1048)")
    txt = txt.replace("integrated Fairlearn's ThresholdOptimizer", "integrated Fairlearn's [6] ThresholdOptimizer")
    txt = txt.replace("reduced the demographic parity difference from 0.105 to 0.000, achieving complete approval rate parity.",
                      "reduced the demographic parity difference from 10.48% (0.1048) to 1.67% (0.0167), achieving near-perfect approval rate parity.")
    txt = txt.replace("resulted in a model accuracy of 76.00% compared to the baseline 73.00%",
                      "resulted in a model accuracy of 75.50% compared to the baseline 73.00%")
    p_321.text = txt
    format_normal_paragraph(p_321)
    print("Updated P321 (Section 8) with citations and metrics.")

    # --- 9. Modify Paragraph 335 (Section 9.3 Future Work) ---
    txt = p_335.text
    txt = txt.replace("real-world lending workflows.", "real-world lending workflows [10].")
    txt = txt.replace("such as Fairlearn's Exponentiated Gradient", "such as Fairlearn's [6] Exponentiated Gradient")
    p_335.text = txt
    format_normal_paragraph(p_335)
    print("Updated P335 (Section 9.3) with citations.")

    # --- 10. Elevate References Heading Style to Heading 1 ---
    p_ref_head.style = 'Heading 1'
    print("Elevated References heading style to Heading 1.")

    # --- 11. Search and Update Table Cells (Verification logs table in Appendix C) ---
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    cell_txt = paragraph.text
                    if "Fairlearn" in cell_txt:
                        cell_txt = cell_txt.replace("via Fairlearn without", "via Fairlearn [6] without")
                        cell_txt = cell_txt.replace("(0.105 for gender)", "(10.48% for gender)")
                        paragraph.text = cell_txt
                        for r in paragraph.runs:
                            r.font.name = 'Times New Roman'
                            r.font.size = Pt(10) # Smaller font for tables

    # --- 12. Delete Old Bibliography Paragraphs Safely ---
    to_delete = []
    found_ref = False
    for p in doc.paragraphs:
        if p._p == p_ref_1._p:
            found_ref = True
        if p._p == p_app_a._p:
            break
        if found_ref:
            to_delete.append(p)
            
    print(f"Deleting {len(to_delete)} old bibliography paragraphs...")
    for p in to_delete:
        p_element = p._p
        parent = p_element.getparent()
        parent.remove(p_element)

    # Re-verify that APPENDIX A is still there in the document structure
    p_app_a_check = None
    for p in doc.paragraphs:
        if p._p == p_app_a._p:
            p_app_a_check = p
            break
            
    if p_app_a_check is None:
        print("Error: APPENDIX A heading was lost during deletion!")
        return

    # --- 13. Insert the 10 New Bibliography References in Sequential Order ---
    refs_data = [
        ("1", 
         "N. Bussmann, P. Giudici, D. Marinelli, and J. Papenbrock, ", 
         "\"Explainable machine learning in credit risk management,\" ", 
         "Computational Economics", 
         ", vol. 57, pp. 203-216, 2021."),
         
        ("2", 
         "M. T. Ribeiro, S. Singh, and C. Guestrin, ", 
         "\"Why should I trust you?: Explaining the predictions of any classifier,\" ", 
         "in Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining (KDD)", 
         ", San Francisco, CA, USA, 2016, pp. 1135–1144."),
         
        ("3", 
         "S. M. Lundberg and S. I. Lee, ", 
         "\"A unified approach to interpreting model predictions,\" ", 
         "in Advances in Neural Information Processing Systems (NeurIPS)", 
         ", Long Beach, CA, USA, 2017, pp. 4765–4774."),
         
        ("4", 
         "C. Agarwal et al., ", 
         "\"OpenXAI: Towards a transparent evaluation of post-hoc explanation methods,\" ", 
         "in Advances in Neural Information Processing Systems (NeurIPS)", 
         ", New Orleans, LA, USA, 2022."),
         
        ("5", 
         "R. K. Mothilal, A. Sharma, and C. Tan, ", 
         "\"Explaining machine learning classifiers through diverse counterfactual explanations,\" ", 
         "in Proceedings of the 2020 Conference on Fairness, Accountability, and Transparency (FAT)", 
         ", Barcelona, Spain, 2020, pp. 607–617."),
         
        ("6", 
         "S. Bird et al., ", 
         "\"Fairlearn: A toolkit for assessing and improving fairness in AI,\" ", 
         "Microsoft Technical Report MSR-TR-2020-32", 
         ", 2020."),
         
        ("7", 
         "UCI Machine Learning Repository, ", 
         "\"Statlog (German Credit Data) Data Set,\" ", 
         "", 
         "2010. [Online]. Available: https://archive.ics.uci.edu/ml/datasets/statlog+(german+credit+data)."),
         
        ("8", 
         "FICO, ", 
         "\"Explainable Machine Learning Challenge: HELOC Dataset,\" ", 
         "", 
         "2018. [Online]. Available: https://community.fico.com/s/explainable-machine-learning-challenge."),
         
        ("9", 
         "S. Barocas, M. Hardt, and A. Narayanan, ", 
         "", 
         "Fairness and Machine Learning: Limitations and Opportunities", 
         ". MIT Press, 2023."),
         
        ("10", 
         "R. Guidotti, A. Monreale, S. Ruggieri, F. Turini, F. Giannotti, and D. Pedreschi, ", 
         "\"A survey of methods for explaining black box models,\" ", 
         "ACM Computing Surveys (CSUR)", 
         ", vol. 51, no. 5, pp. 93:1-93:42, 2018.")
    ]

    for num, auth, tit, jc, rest in refs_data:
        add_reference_entry(doc, p_app_a_check, num, auth, tit, jc, rest)
        
    doc.save(doc_path)
    print("New 10 sequential bibliography references with IEEE italics formatting added successfully!")
    print("Document saved successfully!")

if __name__ == '__main__':
    main()
