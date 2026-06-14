import os
import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def set_table_borders(table):
    tblPr = table._tbl.tblPr
    tblBorders = tblPr.find(qn('w:tblBorders'))
    if tblBorders is None:
        tblBorders = OxmlElement('w:tblBorders')
        tblPr.append(tblBorders)
    else:
        tblBorders.clear()
        
    for border_name, color, sz in [
        ('top', '808080', '4'),
        ('left', '808080', '4'),
        ('bottom', '808080', '4'),
        ('right', '808080', '4'),
        ('insideH', 'C0C0C0', '4'),
        ('insideV', 'C0C0C0', '4')
    ]:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), sz)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tblBorders.append(border)

def format_cell_text(cell, text, bold=False):
    cell.paragraphs[0].text = "" # clear
    p = cell.paragraphs[0]
    p.style = 'Normal'
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    
    # Vertically center
    tcPr = cell._tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), 'center')
    tcPr.append(vAlign)
    
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.bold = bold

def insert_caption_after_paragraph(doc, p, caption_text, table_num):
    caption_p = doc.add_paragraph()
    caption_p.style = 'Caption'
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_p.paragraph_format.space_before = Pt(6)
    caption_p.paragraph_format.space_after = Pt(4)
    
    pPr = caption_p._p.get_or_add_pPr()
    keepNext = OxmlElement('w:keepNext')
    pPr.append(keepNext)
    
    caption_p.add_run("Table ")
    
    fldSimple = OxmlElement('w:fldSimple')
    fldSimple.set(qn('w:instr'), ' SEQ Table \\* ARABIC ')
    r_fld = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    noProof = OxmlElement('w:noProof')
    rPr.append(noProof)
    r_fld.append(rPr)
    t_fld = OxmlElement('w:t')
    t_fld.text = str(table_num)
    r_fld.append(t_fld)
    fldSimple.append(r_fld)
    caption_p._p.append(fldSimple)
    
    caption_p.add_run(f" : {caption_text}")
    
    p_element = p._p
    parent = p_element.getparent()
    parent.insert(parent.index(p_element) + 1, caption_p._p)
    return caption_p

def insert_table_after_paragraph(doc, p, rows, cols, data, headers=None):
    table = doc.add_table(rows, cols)
    p_element = p._p
    parent = p_element.getparent()
    parent.insert(parent.index(p_element) + 1, table._tbl)
    
    table.style = 'Normal Table'
    set_table_borders(table)
    
    # Set Table alignment to center
    tblPr = table._tbl.tblPr
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    tblPr.append(jc)
    
    # Set table width
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is not None:
        tblW.set(qn('w:w'), '9026')
        tblW.set(qn('w:type'), 'dxa')
        
    row_idx = 0
    if headers:
        for col_idx, header in enumerate(headers):
            cell = table.rows[0].cells[col_idx]
            format_cell_text(cell, header, bold=True)
        row_idx = 1
        
    for data_row in data:
        for col_idx, text in enumerate(data_row):
            cell = table.rows[row_idx].cells[col_idx]
            format_cell_text(cell, text, bold=False)
        row_idx += 1
        
    return table

def delete_paragraph(p):
    p_element = p._p
    parent = p_element.getparent()
    parent.remove(p_element)

def main():
    doc_path = r'c:/Users/aliso/Downloads/xai-credit-decision-main/4992FinalReport1011165.docx'
    doc = docx.Document(doc_path)
    
    p_splits = None
    p_hyper = None
    p_conf = None
    p_app_d = None
    p_app_e_placeholder = None
    
    # Locate key paragraphs
    for idx, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text.startswith("The empirical validation of our credit decision framework is based on rigorous data splitting"):
            p_splits = p
            print(f"Found Splits Paragraph at idx {idx}")
        elif text.startswith("The PyTorch Deep Neural Network (DNN) architectures were optimized using grid search"):
            p_hyper = p
            print(f"Found Hyperparameters Paragraph at idx {idx}")
        elif text.startswith("To maximize the F1-score for the German Credit model"):
            p_conf = p
            print(f"Found Confusion Matrix Paragraph at idx {idx}")
        elif text.startswith("APPENDIX D: Meeting Minutes"):
            p_app_d = p
            print(f"Found Appendix D heading at idx {idx}")
        elif text.startswith("[Include the project poster"):
            p_app_e_placeholder = p
            print(f"Found Appendix E placeholder at idx {idx}")
            
    if not (p_splits and p_hyper and p_conf):
        print("Error: Could not locate Appendix C paragraphs.")
        return
        
    # 1. Update text content
    p_splits.text = "The empirical validation of our credit decision framework is based on rigorous data splitting and stratification protocols. To prevent training bias and data leakage, the historical default ratios are preserved identically across all experimental splits using a stratified division. The exact distribution of training, validation, and held-out test samples for both the UCI Statlog German Credit and FICO HELOC datasets is detailed in Table 7."
    # Style runs in p_splits to Times New Roman, 12 pt, and make sure it has style 'Normal'
    p_splits.style = 'Normal'
    for r in p_splits.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        
    p_hyper.text = "The PyTorch Deep Neural Network (DNN) architectures were optimized using grid search hyperparameter tuning. The optimal configurations selected for both datasets, including network depth, layer widths, batch size, learning rate, weight decay, dropout rate, and early stopping patience, are summarized in Table 8."
    p_hyper.style = 'Normal'
    for r in p_hyper.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        
    p_conf.text = "To evaluate the predictive capacity of our models, classification threshold calibration was performed. The calibrated optimal threshold and resulting test metrics are summarized in Table 3 (Section 5.3). The detailed classification outcomes on the held-out test sets are represented as confusion matrices in Table 9 and Table 10, detailing the distribution of true negatives, false positives, false negatives, and true positives for the German Credit and FICO HELOC models, respectively."
    p_conf.style = 'Normal'
    for r in p_conf.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        
    # 2. Insert Table 7 (Dataset Splits) after p_splits
    t7_headers = ['Dataset', 'Total Instances', 'Training Subset (64%)', 'Validation Subset (16%)', 'Test Subset (20%)']
    t7_data = [
        ['UCI Statlog German Credit', '1,000', '640', '160', '200'],
        ['FICO HELOC', '9,872', '6,319', '1,579', '1,974']
    ]
    cap7 = insert_caption_after_paragraph(doc, p_splits, "Dataset Splits and Stratification", 7)
    tbl7 = insert_table_after_paragraph(doc, cap7, 3, 5, t7_data, t7_headers)
    print("Table 7 inserted.")
    
    # 3. Insert Table 8 (Hyperparameters) after p_hyper
    t8_headers = ['Dataset', 'Hidden Layers', 'Batch Size', 'Learning Rate', 'L2 Weight Decay', 'Dropout Rate', 'Early Stopping Patience']
    t8_data = [
        ['German Credit', '64, 32', '64', '0.0005', '0.0001', '0.30', '20 epochs (best: 11)'],
        ['FICO HELOC', '64, 32', '128', '0.0007', '0.0001', '0.25', '12 epochs']
    ]
    cap8 = insert_caption_after_paragraph(doc, p_hyper, "Optimal Hyperparameter Configurations", 8)
    tbl8 = insert_table_after_paragraph(doc, cap8, 3, 7, t8_data, t8_headers)
    print("Table 8 inserted.")
    
    # 4. Insert Table 10 (FICO HELOC Confusion Matrix) and Table 9 (German Credit Confusion Matrix) after p_conf
    t9_headers = ['Actual \\ Predicted', 'Predicted Negative (Bad Credit)', 'Predicted Positive (Good Credit)']
    t9_data = [
        ['Actual Negative (Bad Credit)', '34 (TN)', '26 (FP)'],
        ['Actual Positive (Good Credit)', '28 (FN)', '112 (TP)']
    ]
    
    t10_headers = ['Actual \\ Predicted', 'Predicted Negative', 'Predicted Positive']
    t10_data = [
        ['Actual Negative', '784 (TN)', '244 (FP)'],
        ['Actual Positive', '274 (FN)', '673 (TP)']
    ]
    
    # We want Table 9 to appear before Table 10.
    # To do that, we insert Table 10 first, then insert Table 9 (which pushes Table 10 down), 
    # OR we can just insert Table 9 first, and then insert Table 10 after Table 9.
    # Inserting Table 10 after Table 9 is easier if we insert it relative to the Table 9 elements.
    # Let's do:
    # 1. Insert Table 10 after p_conf
    # 2. Insert Table 9 after p_conf
    # This results in: p_conf -> Table 9 -> Table 10.
    # Let's write the sequence:
    cap10 = insert_caption_after_paragraph(doc, p_conf, "Confusion Matrix for FICO HELOC Model", 10)
    tbl10 = insert_table_after_paragraph(doc, cap10, 3, 3, t10_data, t10_headers)
    
    cap9 = insert_caption_after_paragraph(doc, p_conf, "Confusion Matrix for German Credit Model", 9)
    tbl9 = insert_table_after_paragraph(doc, cap9, 3, 3, t9_data, t9_headers)
    print("Tables 9 and 10 inserted.")
    
    # 5. Clean up the empty paragraphs between the end of Appendix C (now after tbl10) and Appendix D
    # We find all empty paragraphs between p_conf and p_app_d and delete them.
    # We can search through the document paragraphs list.
    to_delete = []
    found_conf = False
    for p in doc.paragraphs:
        if p == p_conf:
            found_conf = True
            continue
        if p == p_app_d:
            break
        if found_conf:
            if p.text.strip() == "":
                to_delete.append(p)
                
    for p in to_delete:
        try:
            delete_paragraph(p)
        except Exception as e:
            print(f"Could not delete empty paragraph: {e}")
    print(f"Deleted {len(to_delete)} empty paragraphs in Appendix C.")
    
    # 6. Clear Appendix E placeholder
    if p_app_e_placeholder:
        p_app_e_placeholder.text = ""
        p_app_e_placeholder.style = 'Normal'
        print("Cleared Appendix E placeholder.")
        
    # Save the document
    doc.save(doc_path)
    print("Document saved successfully!")

if __name__ == '__main__':
    main()
